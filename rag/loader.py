"""
Document Loader — Extracts text from PDF, TXT, MD, and DOCX files.

Supports Streamlit UploadedFile objects and local file paths.
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import BinaryIO

from config import SUPPORTED_EXTENSIONS


@dataclass
class Document:
    """A loaded document with its text content and metadata."""

    text: str
    filename: str
    file_type: str
    num_characters: int = field(init=False)
    num_words: int = field(init=False)

    def __post_init__(self):
        self.num_characters = len(self.text)
        self.num_words = len(self.text.split())

    def __repr__(self) -> str:
        return (
            f"Document(filename={self.filename!r}, "
            f"type={self.file_type!r}, "
            f"words={self.num_words:,})"
        )


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyPDF2."""
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_text(file_bytes: bytes) -> str:
    """Decode raw bytes as UTF-8 text."""
    return file_bytes.decode("utf-8", errors="replace")


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".txt": _extract_text,
    ".md": _extract_text,
    ".docx": _extract_docx,
}


def load_document(source: str | Path | BinaryIO, filename: str | None = None) -> Document:
    """
    Load a document from a file path or file-like object.

    Parameters
    ----------
    source : str, Path, or file-like
        A local file path or a file-like object (e.g. Streamlit UploadedFile).
    filename : str, optional
        Override filename (useful when source is a file-like object).

    Returns
    -------
    Document
        Parsed document with text content and metadata.

    Raises
    ------
    ValueError
        If the file type is not supported.
    """
    # Resolve filename and bytes
    if isinstance(source, (str, Path)):
        path = Path(source)
        filename = filename or path.name
        file_bytes = path.read_bytes()
    else:
        # File-like object (Streamlit UploadedFile)
        filename = filename or getattr(source, "name", "unknown")
        file_bytes = source.read()
        if hasattr(source, "seek"):
            source.seek(0)

    ext = Path(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {supported}"
        )

    extractor = _EXTRACTORS[ext]
    text = extractor(file_bytes)

    if not text.strip():
        raise ValueError(f"No text content could be extracted from '{filename}'.")

    return Document(text=text, filename=filename, file_type=ext)


def load_multiple(sources: list) -> list[Document]:
    """Load multiple documents, skipping failures with warnings."""
    documents = []
    errors = []
    for src in sources:
        try:
            fname = getattr(src, "name", None)
            documents.append(load_document(src, filename=fname))
        except (ValueError, Exception) as e:
            name = getattr(src, "name", str(src))
            errors.append(f"{name}: {e}")
    return documents, errors
